import streamlit as st
import requests
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import urllib.parse

st.title("Dashboard Laporan Rekap Publikasi Media Online")

API_KEY = os.getenv("APIFLASH_KEY")

if not API_KEY:
    st.error("API Key belum diisi di Secrets.")
    st.stop()

links_input = st.text_area("Masukkan link (1 link per baris)")

if st.button("Buat Laporan"):

    if not links_input.strip():
        st.warning("Masukkan link terlebih dahulu.")
        st.stop()

    # Bersihkan link
    links = [l.strip() for l in links_input.split("\n") if l.strip()]

    document = Document()
    document.add_heading("LAPORAN REKAP PUBLIKASI MEDIA ONLINE", level=1)

    table = document.add_table(rows=2, cols=6)
    table.style = "Table Grid"

    headers = ["NO", "TANGGAL", "JUDUL KEGIATAN", "LINK PUBLIKASI", "DOKUMENTASI", "KET"]
    for i in range(6):
        table.rows[0].cells[i].text = headers[i]

    row = table.rows[1]
    row.cells[0].text = "1"
    row.cells[1].text = datetime.now().strftime("%d/%m/%Y")

    # =========================
    # AMBIL JUDUL DARI SLUG URL
    # =========================

    judul_list = []

    for link in links:
        try:
            slug = link.rstrip("/").split("/")[-1]
            slug = slug.replace("-", " ")
            title = slug.title()
            judul_list.append(title)
        except:
            judul_list.append("Judul tidak ditemukan")

    row.cells[2].text = "\n".join(
        [f"{i+1}. {judul}" for i, judul in enumerate(judul_list)]
    )

    row.cells[3].text = "\n".join(
        [f"{i+1}. {link}" for i, link in enumerate(links)]
    )

    # =========================
    # SCREENSHOT
    # =========================

    for i, link in enumerate(links, start=1):
        try:
            encoded_url = urllib.parse.quote(link, safe="")

            screenshot_url = (
                f"https://api.apiflash.com/v1/urltoimage?"
                f"access_key={API_KEY}"
                f"&url={encoded_url}"
                f"&width=1280"
                f"&height=800"
                f"&format=png"
                f"&fresh=true"
            )

            response = requests.get(screenshot_url, timeout=30)

            if response.status_code == 200 and "image" in response.headers.get("Content-Type", ""):
                image_path = f"screenshot_{i}.png"

                with open(image_path, "wb") as f:
                    f.write(response.content)

                row.cells[4].paragraphs[0].add_run().add_picture(
                    image_path, width=Inches(1.2)
                )
            else:
                row.cells[4].paragraphs[0].add_run(
                    f"\nScreenshot gagal untuk link {i}\n"
                )

        except:
            row.cells[4].paragraphs[0].add_run(
                f"\nError pada link {i}\n"
            )

    row.cells[5].text = "-"

    file_name = "Laporan_Rekap_Publikasi.docx"
    document.save(file_name)

    with open(file_name, "rb") as f:
        st.download_button(
            "Download Laporan Word",
            f,
            file_name=file_name
        )
